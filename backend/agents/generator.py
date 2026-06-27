import os
from docx import Document
from backend.models.schemas import CandidateProfile

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), '..', 'templates', 'interlev_template.docx')
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', '..', 'output')

def ensure_template_exists():
    """Create a basic placeholder template if it doesn't exist."""
    os.makedirs(os.path.dirname(TEMPLATE_PATH), exist_ok=True)
    if not os.path.exists(TEMPLATE_PATH):
        doc = Document()
        doc.add_heading('INTERLEV Freelancer Profile', 0)
        doc.add_paragraph('Name: {{NAME}}')
        doc.add_paragraph('Role: {{ROLE}}')
        doc.add_paragraph('Experience: {{EXPERIENCE}} years')
        doc.add_heading('Summary', level=1)
        doc.add_paragraph('{{SUMMARY}}')
        doc.add_heading('Skills', level=1)
        doc.add_paragraph('{{SKILLS}}')
        doc.add_heading('Languages', level=1)
        doc.add_paragraph('{{LANGUAGES}}')
        doc.save(TEMPLATE_PATH)

def generate_interlev_cv(profile: CandidateProfile) -> str:
    """
    Generates a tailored Word document based on the INTERLEV template.
    Returns the path to the generated file.
    """
    ensure_template_exists()
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    doc = Document(TEMPLATE_PATH)
    
    # Dictionary of replacements
    replacements = {
        '{{NAME}}': profile.name or "N/A",
        '{{ROLE}}': profile.role or "N/A",
        '{{EXPERIENCE}}': str(profile.experience_years) if profile.experience_years else "N/A",
        '{{SUMMARY}}': profile.summary or "N/A",
        '{{SKILLS}}': ", ".join(profile.skills) if profile.skills else "N/A",
        '{{LANGUAGES}}': ", ".join(profile.languages) if profile.languages else "N/A"
    }
    
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)

    output_filename = f"INTERLEV_CV_{profile.name.replace(' ', '_') if profile.name else 'Candidate'}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    doc.save(output_path)
    
    # We could convert to PDF here using docx2pdf
    # For MVP, we will just return the DOCX path, since docx2pdf requires MS Word installed on Windows
    # and LibreOffice is complex to automate without knowing the exact installation path.
    
    return output_path
