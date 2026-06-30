import os
import re
import tempfile
import textwrap
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document

from backend.app.models.candidate import Candidate
from backend.app.models.job import Job
from backend.app.services.app_settings import load_settings

class CVFormatterAgent:
    def __init__(self, output_dir: str = "formatted_cvs"):
        if os.getenv("VERCEL") and output_dir == "formatted_cvs":
            output_dir = str(
                Path(tempfile.gettempdir()) / "interlev-agent" / "formatted_cvs"
            )
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_file(
        self,
        candidate: Candidate,
        job: Optional[Job] = None,
        output_format: Optional[str] = None,
        content_overrides: Optional[Dict[str, Any]] = None,
    ) -> str:
        selected_format = (output_format or load_settings().cv_format.output_format or "docx").lower()
        if selected_format == "pdf":
            return self.generate_pdf(candidate, job, content_overrides)
        return self.generate_docx(candidate, job.title if job else None, job, content_overrides)

    def build_preview(self, candidate: Candidate, job: Optional[Job] = None) -> Dict[str, Any]:
        settings = load_settings()
        return {
            "candidate_id": candidate.id,
            "job_id": job.id if job else None,
            "template_name": settings.cv_format.template_name,
            "template_file_name": settings.cv_format.template_file_name,
            "output_format": settings.cv_format.output_format,
            "file_name": self._file_name(candidate, job, settings.cv_format.output_format),
            "sections": self._cv_sections(candidate, job),
        }

    def generate_docx(
        self,
        candidate: Candidate,
        job_title: Optional[str] = None,
        job: Optional[Job] = None,
        content_overrides: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        Generates an INTERLEV formatted CV in DOCX format.
        """
        settings = load_settings()
        template_path = settings.cv_format.template_file_path
        use_uploaded_template = (
            template_path
            and template_path.lower().endswith(".docx")
            and os.path.exists(template_path)
        )
        doc = Document(template_path) if use_uploaded_template else Document()

        replaced = False
        if use_uploaded_template:
            replaced = self._replace_placeholders(doc, candidate, job, job_title, content_overrides)
            if not replaced:
                replaced = self._fill_uploaded_docx_template(
                    doc,
                    candidate,
                    job,
                    job_title,
                    content_overrides,
                )
        else:
            title = doc.add_heading(settings.cv_format.template_name, 0)
            title.alignment = 1 # Center

        if not replaced:
            if use_uploaded_template:
                doc.add_page_break()
            for section in self._cv_sections(candidate, job, job_title, content_overrides):
                doc.add_heading(f"{section['title']}:", level=1)
                doc.add_paragraph(section["body"])

        # Save the document
        filename = self._file_name(candidate, job, "docx")
        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)

        return file_path

    def generate_pdf(
        self,
        candidate: Candidate,
        job: Optional[Job] = None,
        content_overrides: Optional[Dict[str, Any]] = None,
    ) -> str:
        settings = load_settings()
        template_path = settings.cv_format.template_file_path
        if (
            template_path
            and template_path.lower().endswith(".pdf")
            and os.path.exists(template_path)
        ):
            return self._generate_pdf_from_uploaded_template(candidate, job, template_path, content_overrides)

        try:
            import fitz
        except Exception as exc:
            raise RuntimeError("PDF generation requires PyMuPDF.") from exc

        file_path = os.path.join(self.output_dir, self._file_name(candidate, job, "pdf"))
        doc = fitz.open()
        page = doc.new_page(width=595, height=842)
        margin = 52
        y = 54

        def add_page() -> None:
            nonlocal page, y
            page = doc.new_page(width=595, height=842)
            y = 54

        def write_lines(text: str, size: int = 11, color=(0, 0, 0), line_height: int = 15) -> None:
            nonlocal y
            lines: List[str] = []
            for paragraph in str(text or "").splitlines() or [""]:
                wrapped = textwrap.wrap(paragraph, width=88) or [""]
                lines.extend(wrapped)
            for line in lines:
                if y > 790:
                    add_page()
                page.insert_text((margin, y), line, fontsize=size, fontname="helv", color=color)
                y += line_height

        page.draw_rect(fitz.Rect(0, 0, 595, 112), color=(0.06, 0.09, 0.18), fill=(0.06, 0.09, 0.18))
        page.insert_text((margin, 50), settings.cv_format.template_name, fontsize=19, fontname="helv", color=(1, 1, 1))
        page.insert_text((margin, 78), candidate.name or "Candidate", fontsize=15, fontname="helv", color=(0.72, 0.9, 1))
        page.insert_text(
            (margin, 98),
            self._pdf_subtitle(candidate, job, settings.cv_format.template_file_name),
            fontsize=11,
            fontname="helv",
            color=(0.85, 0.9, 0.98),
        )
        y = 142

        for section in self._cv_sections(candidate, job, content_overrides=content_overrides):
            if y > 760:
                add_page()
            page.insert_text((margin, y), section["title"].upper(), fontsize=12, fontname="helv", color=(0.05, 0.23, 0.55))
            y += 18
            write_lines(section["body"], size=10, color=(0.1, 0.13, 0.2), line_height=14)
            y += 12

        page.insert_text(
            (margin, 820),
            f"Generated by INTERLEV on {datetime.utcnow().strftime('%Y-%m-%d')}",
            fontsize=8,
            fontname="helv",
            color=(0.45, 0.48, 0.55),
        )
        doc.save(file_path)
        doc.close()
        return file_path

    def _generate_pdf_from_uploaded_template(
        self,
        candidate: Candidate,
        job: Optional[Job],
        template_path: str,
        content_overrides: Optional[Dict[str, Any]] = None,
    ) -> str:
        try:
            import fitz
        except Exception as exc:
            raise RuntimeError("PDF generation requires PyMuPDF.") from exc

        file_path = os.path.join(self.output_dir, self._file_name(candidate, job, "pdf"))
        doc = fitz.open(template_path)
        values = self._pdf_placeholder_values(candidate, job, content_overrides)

        for page in doc:
            had_known_placeholders = self._template_has_known_placeholders(page)
            for placeholder, value in values.items():
                self._replace_pdf_text(page, placeholder, value)
            self._replace_pdf_template_blocks(page, candidate, job, content_overrides)

            if not had_known_placeholders:
                self._write_sections_on_pdf_template(page, candidate, job, content_overrides)

        doc.save(file_path, garbage=4, deflate=True)
        doc.close()
        return file_path

    def _replace_pdf_text(self, page, placeholder: str, value: str) -> bool:
        rects = page.search_for(placeholder)
        if not rects:
            return False

        for rect in rects:
            sidebar = self._is_pdf_sidebar_text(page, rect)
            padding = 1.5
            target = rect + (-padding, -padding, padding, padding)
            page.add_redact_annot(
                target,
                fill=(0.06, 0.09, 0.16) if sidebar else (1, 1, 1),
            )
        page.apply_redactions()

        for rect in rects:
            sidebar = self._is_pdf_sidebar_text(page, rect)
            fontsize = max(7, min(20, rect.height * 0.72))
            if sidebar:
                fontsize = min(fontsize, 8)
                target = (rect.x0, rect.y0 - 1, min(page.rect.width * 0.33, rect.x0 + 128), rect.y1 + 7)
            else:
                target = rect + (0, -1, max(30, rect.width * 2.0), max(8, rect.height * 1.6))
            page.insert_textbox(
                target,
                str(value or ""),
                fontsize=fontsize,
                fontname="helv",
                color=(1, 1, 1) if sidebar else (0.08, 0.09, 0.12),
                align=0,
            )
        return True

    def _is_pdf_sidebar_text(self, page, rect) -> bool:
        return rect.x0 < float(page.rect.width) * 0.32 and rect.y0 > 120

    def _replace_pdf_template_blocks(
        self,
        page,
        candidate: Candidate,
        job: Optional[Job],
        content_overrides: Optional[Dict[str, Any]] = None,
    ) -> None:
        phone_rects = page.search_for("Phone:")
        if phone_rects:
            rect = phone_rects[0]
            block = (rect.x0, rect.y0 - 2, page.rect.width - 42, rect.y1 + 5)
            page.add_redact_annot(block, fill=(1, 1, 1))
            page.apply_redactions()
            contact_parts = []
            phone = self._override_value(content_overrides, "phone", candidate.phone or "")
            email = self._override_value(content_overrides, "email", candidate.email or "")
            location = self._override_value(content_overrides, "location", candidate.location or "")
            if phone:
                contact_parts.append(f"Phone: {phone}")
            if email:
                contact_parts.append(f"Email: {email}")
            if location:
                contact_parts.append(f"Location: {location}")
            page.insert_textbox(
                block,
                "  |  ".join(contact_parts) or "Contact details available upon request",
                fontsize=8.3,
                fontname="helv",
                color=(0.08, 0.09, 0.12),
                align=0,
            )

        summary_anchor = "[Write 3-4 lines about your experience, main skills, achievements, and the type of"
        rects = page.search_for(summary_anchor)
        if not rects:
            return

        rect = rects[0]
        block = (rect.x0, rect.y0 - 2, page.rect.width - 42, rect.y0 + 58)
        page.add_redact_annot(block, fill=(1, 1, 1))
        page.apply_redactions()
        summary = self._compact_text_for_pdf_box(
            self._override_value(content_overrides, "summary", candidate.summary or "N/A"),
            width=74,
            max_lines=5,
        )
        page.insert_textbox(
            block,
            summary,
            fontsize=7.2,
            fontname="helv",
            color=(0.08, 0.09, 0.12),
            align=0,
        )

    def _compact_text_for_pdf_box(self, text: str, width: int, max_lines: int) -> str:
        clean = re.sub(r"\s+", " ", str(text or "")).strip()
        lines = textwrap.wrap(clean, width=width) or [""]
        truncated = lines[:max_lines]
        if len(lines) > max_lines and truncated:
            truncated[-1] = truncated[-1].rstrip(" .,;:") + "..."
        return "\n".join(truncated)

    def _template_has_known_placeholders(self, page) -> bool:
        return bool(page.search_for("[YOUR FULL NAME]") or page.search_for("{{NAME}}"))

    def _write_sections_on_pdf_template(
        self,
        page,
        candidate: Candidate,
        job: Optional[Job],
        content_overrides: Optional[Dict[str, Any]] = None,
    ) -> None:
        width = float(page.rect.width)
        margin = 42
        y = 50
        max_width = width - (margin * 2)

        for section in self._cv_sections(candidate, job, content_overrides=content_overrides):
            if y > float(page.rect.height) - 80:
                break
            page.insert_textbox(
                (margin, y, margin + max_width, y + 18),
                section["title"].upper(),
                fontsize=10,
                fontname="helv",
                color=(0.05, 0.23, 0.55),
            )
            y += 18
            lines = textwrap.wrap(str(section["body"] or ""), width=88) or [""]
            text = "\n".join(lines[:8])
            block_height = min(104, 14 * (len(lines[:8]) + 1))
            page.insert_textbox(
                (margin, y, margin + max_width, y + block_height),
                text,
                fontsize=9,
                fontname="helv",
                color=(0.1, 0.13, 0.2),
            )
            y += block_height + 10

    def _pdf_placeholder_values(
        self,
        candidate: Candidate,
        job: Optional[Job],
        content_overrides: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, str]:
        values = self._placeholder_values(candidate, job, None)
        role = self._override_value(content_overrides, "role", values["{{ROLE}}"])
        name = self._override_value(content_overrides, "name", candidate.name or "N/A")
        summary = self._override_value(content_overrides, "summary", candidate.summary or "N/A")
        required = ", ".join(job.required_skills or []) if job else ""
        target = self._override_value(content_overrides, "target_job", job.title if job else role)
        skills = self._split_edit_lines(
            self._override_value(
                content_overrides,
                "skills",
                ", ".join([skill.skill_name for skill in candidate.skills if skill.skill_name]),
            )
        )
        tools = self._split_edit_lines(
            self._override_value(content_overrides, "tools", ", ".join(skills[:3]))
        )
        languages = self._split_edit_lines(
            self._override_value(content_overrides, "languages", "English")
        )
        certifications = self._split_edit_lines(
            self._override_value(content_overrides, "certifications", "Available upon request")
        )
        work = self._split_edit_lines(
            self._override_value(
                content_overrides,
                "work_experience",
                (
                    f"{role}\n"
                    f"{candidate.location or 'Remote'}\n"
                    "Relevant experience tailored for the target role.\n"
                    f"{', '.join(skills[:6]) or 'N/A'}"
                ),
            )
        )
        education = self._split_edit_lines(
            self._override_value(content_overrides, "education", "Available upon request")
        )
        projects = self._split_edit_lines(
            self._override_value(
                content_overrides,
                "projects",
                f"{target or role}\nPrepared for {target or role}. Required skills: {required or 'Not listed'}.",
            )
        )
        references = self._split_edit_lines(
            self._override_value(content_overrides, "references", "Available upon request.")
        )

        pdf_values: Dict[str, str] = {
            "{{NAME}}": name,
            "{{ROLE}}": role,
            "{{SUMMARY}}": summary,
            "{{SKILLS}}": ", ".join(skills) or "N/A",
            "{{TARGET_JOB}}": target or role,
            "{{JOB_COMPANY}}": job.company if job else "",
            "{{JOB_PLATFORM}}": job.platform if job else "",
            "{{REQUIRED_SKILLS}}": required or "Not listed",
            "[YOUR FULL NAME]": name,
            "[Job Title - e.g., Software Engineer / Accountant / Marketing Executive]": role,
            "[Job Title] - [Company Name]": self._line_or(work, 0, role),
            "[Month Year - Month Year] | [City / Remote]": self._line_or(work, 1, candidate.location or "Remote"),
            "[Achievement or responsibility with numbers/results if possible]": self._line_or(
                work,
                2,
                "Relevant experience tailored for the target role.",
            ),
            "[Achievement or responsibility]": self._line_or(work, 3, f"Target opportunity: {target}"),
            "[Tools/skills used]": self._line_or(work, 4, ", ".join(skills[:6]) or "N/A"),
            "[Degree / Diploma / Course Name]": self._line_or(education, 0, "Available upon request"),
            "[Institution Name] | [Year]": "",
            "[Relevant subjects, grade, or achievements]": self._line_or(education, 1, ""),
            "[Project Name]": self._line_or(projects, 0, target or "Relevant project"),
            "[Short project description, your role, tools used, and result/link]": self._line_or(
                projects,
                1,
                f"Prepared for {target}. Required skills: {required or 'Not listed'}.",
            ),
            "[Certificate Name]": self._line_or(certifications, 0, "Available upon request"),
            "[Tool / Software 1]": self._line_or(tools, 0, ""),
            "[Tool / Software 2]": self._line_or(tools, 1, ""),
            "[Tool / Software 3]": self._line_or(tools, 2, ""),
            "[Language 1 - Level]": self._line_or(languages, 0, "English"),
            "[Language 2 - Level]": self._line_or(languages, 1, ""),
            "Available upon request.": self._line_or(references, 0, "Available upon request."),
        }
        for index in range(1, 7):
            pdf_values[f"[Skill {index}]"] = skills[index - 1] if len(skills) >= index else ""
        return pdf_values

    def _override_value(
        self,
        content_overrides: Optional[Dict[str, Any]],
        key: str,
        default: str,
    ) -> str:
        if not content_overrides:
            return str(default or "")
        value = content_overrides.get(key)
        if value is None:
            return str(default or "")
        clean_value = str(value).strip()
        return clean_value if clean_value else str(default or "")

    def _split_edit_lines(self, value: str) -> List[str]:
        parts = re.split(r"[\n,;]+", str(value or ""))
        return [part.strip(" -\t") for part in parts if part.strip(" -\t")]

    def _line_or(self, lines: List[str], index: int, default: str) -> str:
        return lines[index] if len(lines) > index and lines[index] else default

    def _convert_to_pdf(self, docx_path: str) -> str:
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
        try:
            from docx2pdf import convert

            convert(docx_path, pdf_path)
            if os.path.exists(pdf_path):
                return pdf_path
        except Exception:
            pass

        return docx_path

    def _cv_sections(
        self,
        candidate: Candidate,
        job: Optional[Job] = None,
        job_title: Optional[str] = None,
        content_overrides: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, str]]:
        skills_text = self._override_value(
            content_overrides,
            "skills",
            ", ".join([skill.skill_name for skill in candidate.skills]) or "N/A",
        )
        target_role = self._override_value(
            content_overrides,
            "role",
            job_title or (job.title if job else None) or candidate.main_role or "N/A",
        )
        sections = [
            {"title": "Name", "body": self._override_value(content_overrides, "name", candidate.name or "N/A")},
            {"title": "Role", "body": target_role},
            {
                "title": "Professional Summary",
                "body": self._override_value(content_overrides, "summary", candidate.summary or "N/A"),
            },
            {"title": "Key Skills", "body": skills_text},
            {
                "title": "Availability",
                "body": self._override_value(content_overrides, "availability", candidate.availability or "Immediate"),
            },
            {
                "title": "Expected Rate",
                "body": self._override_value(content_overrides, "expected_rate", candidate.expected_rate or "TBD"),
            },
        ]
        if job:
            required = ", ".join(job.required_skills or []) or "Not listed"
            sections.insert(
                2,
                {
                    "title": "Target Opportunity",
                    "body": self._override_value(
                        content_overrides,
                        "target_opportunity",
                        (
                            f"{job.title or 'Job'} at {job.company or job.platform or 'client'}."
                            f" Platform: {job.platform or 'N/A'}. Required skills: {required}."
                        ),
                    ),
                },
            )
        return sections

    def _file_name(self, candidate: Candidate, job: Optional[Job], extension: str) -> str:
        name = self._safe_part(candidate.name or "Candidate")
        role = self._safe_part((job.title if job else candidate.main_role) or "CV")
        return f"INTERLEV_{name}_{role}_{candidate.id}.{extension.lower()}"

    def _safe_part(self, value: str) -> str:
        safe = re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_")
        return safe[:48] or "Profile"

    def _replace_placeholders(
        self,
        doc: Document,
        candidate: Candidate,
        job: Optional[Job],
        job_title: Optional[str],
        content_overrides: Optional[Dict[str, Any]] = None,
    ) -> bool:
        values = self._placeholder_values(candidate, job, job_title, content_overrides)
        replaced = False

        def replace_in_paragraph(paragraph) -> None:
            nonlocal replaced
            text = paragraph.text
            updated = text
            for key, value in values.items():
                if key in updated:
                    updated = updated.replace(key, value)
            if updated != text:
                paragraph.text = updated
                replaced = True

        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)
        return replaced

    def _fill_uploaded_docx_template(
        self,
        doc: Document,
        candidate: Candidate,
        job: Optional[Job],
        job_title: Optional[str],
        content_overrides: Optional[Dict[str, Any]] = None,
    ) -> bool:
        values = self._placeholder_values(candidate, job, job_title, content_overrides)
        sections = {
            section["title"].lower(): section["body"]
            for section in self._cv_sections(candidate, job, job_title, content_overrides)
        }
        skills = self._split_edit_lines(values.get("{{SKILLS}}", ""))
        required = self._split_edit_lines(values.get("{{REQUIRED_SKILLS}}", ""))
        tools = self._split_edit_lines(self._override_value(content_overrides, "tools", values.get("{{SKILLS}}", "")))
        certifications = self._split_edit_lines(
            self._override_value(content_overrides, "certifications", "Available upon request")
        )
        languages = self._split_edit_lines(self._override_value(content_overrides, "languages", "English"))
        work_location = candidate.location or (job.location if job else "") or "Remote"
        work_lines = self._split_edit_lines(
            self._override_value(
                content_overrides,
                "work_experience",
                (
                    f"{values.get('{{ROLE}}', '')}\n"
                    f"{work_location}\n"
                    f"{sections.get('professional summary', '')}\n"
                    f"{values.get('{{SKILLS}}', '')}"
                ),
            )
        )
        project_lines = self._split_edit_lines(
            self._override_value(
                content_overrides,
                "projects",
                (
                    f"{values.get('{{TARGET_JOB}}', values.get('{{ROLE}}', ''))}\n"
                    f"{sections.get('target opportunity', '')}"
                ),
            )
        )

        role = values.get("{{ROLE}}", "Professional Profile")
        name = values.get("{{NAME}}", candidate.name or "Candidate")
        target = values.get("{{TARGET_JOB}}", role)
        platform = values.get("{{JOB_PLATFORM}}", "") or (job.platform if job else "") or "Client"
        company = values.get("{{JOB_COMPANY}}", "") or (job.company if job else "") or platform
        location = values.get("{{LOCATION}}", "") or (job.location if job else "") or candidate.location or "Remote"
        summary = values.get("{{SUMMARY}}", "") or sections.get("professional summary", "")
        target_opportunity = sections.get("target opportunity", "")
        skill_text = ", ".join(skills[:8]) or "N/A"
        required_text = ", ".join(required[:8]) or "Not listed"
        seniority = candidate.experience_level or (
            f"{candidate.experience_years} years" if candidate.experience_years else "Profile-based"
        )

        key_facts = [
            f"Matched for {target} with {required_text}.",
            f"Strong fit for {platform or company} based on skills: {skill_text}.",
            f"Seniority: {seniority}; availability: {candidate.availability or 'Immediate'}.",
            f"Location/work mode: {location}.",
            f"Delivery focus for {company or 'the client'} with job-specific CV tailoring.",
        ]
        industry_fit = [
            f"Relevant profile for {company or platform} and the target role {target}.",
            f"Core skills mapped to this opportunity: {skill_text}.",
            f"Can support the listed requirements: {required_text}.",
        ]
        value_points = [
            f"Prepared specifically for {target_opportunity or target}.",
            f"Highlights the candidate data that best matches the job description.",
            f"Keeps the uploaded INTERLEV template structure while replacing sample content.",
            f"Final document is generated from the active Settings CV template.",
        ]

        changed = False
        key_fact_index = 0
        industry_index = 0
        value_index = 0
        cert_index = 0
        in_industry_section = False
        in_value_section = False

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            upper = text.upper()
            if upper == "BRANCHENWISSEN UND ARBEITSWEISE":
                in_industry_section = True
                in_value_section = False
                continue
            if upper == "KONKRETE MEHRWERTE FÜR DAS PROJEKT":
                in_industry_section = False
                in_value_section = True
                continue
            if upper in {"EXECUTIVE SUMMARY", "KEY FACTS", "KERNKOMPETENZEN", "PROJEKTHISTORIE", "ZERTIFIZIERUNGEN & WEITERBILDUNG"}:
                in_industry_section = False
                in_value_section = False
                continue

            if text.startswith("Kurze professionelle Zusammenfassung"):
                self._set_paragraph_text(
                    paragraph,
                    self._compact_text(
                        summary or f"{name} is prepared for {target}. {target_opportunity}",
                        620,
                    ),
                )
                changed = True
            elif text.startswith("• Relevante") or text.startswith("• Technologische") or text.startswith("• Erfahrung") or text.startswith("• Kommunikations") or text.startswith("• Schnelle"):
                self._set_paragraph_text(paragraph, f"• {key_facts[min(key_fact_index, len(key_facts) - 1)]}")
                key_fact_index += 1
                changed = True
            elif text.startswith("• PMP") or text.startswith("• Microsoft") or text.startswith("• Fachliche"):
                self._set_paragraph_text(
                    paragraph,
                    f"• {self._line_or(certifications, cert_index, 'Available upon request')}",
                )
                cert_index += 1
                changed = True
            elif in_industry_section and text.startswith("•"):
                self._set_paragraph_text(paragraph, f"• {industry_fit[min(industry_index, len(industry_fit) - 1)]}")
                industry_index += 1
                changed = True
            elif in_value_section and text.startswith("•"):
                self._set_paragraph_text(paragraph, f"• {value_points[min(value_index, len(value_points) - 1)]}")
                value_index += 1
                changed = True

        if doc.tables:
            self._fill_header_table(doc.tables[0], candidate, role, seniority, name)
            changed = True
        if len(doc.tables) > 1:
            self._fill_focus_table(doc.tables[1], skills, required, target, platform, languages)
            changed = True
        if len(doc.tables) > 2:
            self._fill_competency_table(doc.tables[2], skills, required, tools, target)
            changed = True
        for index, table in enumerate(doc.tables[3:6]):
            self._fill_project_table(
                table,
                index,
                candidate,
                job,
                role,
                target,
                company,
                platform,
                location,
                skill_text,
                required_text,
                summary,
                target_opportunity,
                work_lines,
                project_lines,
            )
            changed = True

        return changed

    def _set_paragraph_text(self, paragraph, text: str) -> None:
        paragraph.text = str(text or "")

    def _set_cell_text(self, cell, text: str) -> None:
        cell.text = str(text or "")

    def _fill_header_table(
        self,
        table,
        candidate: Candidate,
        role: str,
        seniority: str,
        name: str,
    ) -> None:
        if not table.rows or len(table.rows[0].cells) < 1:
            return
        self._set_cell_text(
            table.rows[0].cells[0],
            "\n".join(
                [
                    f"Kandidaten-ID: INT-KAN-{int(candidate.id or 0):04d}",
                    f"Name: {name}",
                    f"Zielrolle: {role}",
                    f"Seniorität: {seniority}",
                ]
            ),
        )

    def _fill_focus_table(
        self,
        table,
        skills: List[str],
        required: List[str],
        target: str,
        platform: str,
        languages: List[str],
    ) -> None:
        if len(table.rows) < 2 or len(table.rows[1].cells) < 3:
            return
        row = table.rows[1].cells
        self._set_cell_text(row[0], ", ".join(skills[:4]) or target)
        self._set_cell_text(row[1], platform or "Remote / freelance")
        self._set_cell_text(
            row[2],
            f"Job-specific fit: {', '.join(required[:4]) or 'requirements mapped'}; Languages: {', '.join(languages[:2]) or 'English'}",
        )

    def _fill_competency_table(
        self,
        table,
        skills: List[str],
        required: List[str],
        tools: List[str],
        target: str,
    ) -> None:
        rows = [
            ("Core skills", ", ".join(skills[:6]) or "N/A"),
            ("Required match", ", ".join(required[:6]) or "Not listed"),
            ("Tools / software", ", ".join(tools[:5]) or ", ".join(skills[:5]) or "N/A"),
            ("Delivery focus", f"Tailored for {target}"),
            ("Communication", "Clear client communication, fast onboarding, documented delivery"),
        ]
        for row_index, (area, description) in enumerate(rows, start=1):
            if len(table.rows) <= row_index or len(table.rows[row_index].cells) < 2:
                break
            self._set_cell_text(table.rows[row_index].cells[0], area)
            self._set_cell_text(table.rows[row_index].cells[1], description)

    def _fill_project_table(
        self,
        table,
        index: int,
        candidate: Candidate,
        job: Optional[Job],
        role: str,
        target: str,
        company: str,
        platform: str,
        location: str,
        skill_text: str,
        required_text: str,
        summary: str,
        target_opportunity: str,
        work_lines: List[str],
        project_lines: List[str],
    ) -> None:
        period = ["Target project", "Relevant profile", "Availability"][min(index, 2)]
        client = [company or platform, platform or company, location][min(index, 2)]
        project_role = [target, self._line_or(work_lines, 0, role), role][min(index, 2)]
        technology = [required_text, skill_text, skill_text][min(index, 2)]
        description = [
            target_opportunity or self._line_or(project_lines, 1, f"Prepared for {target}."),
            summary or self._line_or(work_lines, 2, "Relevant experience mapped from candidate CV."),
            f"Available for {job.contract_type if job and job.contract_type else 'remote/freelance'} work. Expected rate: {candidate.expected_rate or 'TBD'}.",
        ][min(index, 2)]
        values = [period, client, project_role, technology, self._compact_text(description, 420)]
        for row_index, value in enumerate(values):
            if len(table.rows) <= row_index or len(table.rows[row_index].cells) < 2:
                break
            self._set_cell_text(table.rows[row_index].cells[1], value)

    def _compact_text(self, text: str, max_chars: int) -> str:
        clean = re.sub(r"\s+", " ", str(text or "")).strip()
        if len(clean) <= max_chars:
            return clean
        return clean[: max_chars - 3].rstrip(" .,;:") + "..."

    def _placeholder_values(
        self,
        candidate: Candidate,
        job: Optional[Job],
        job_title: Optional[str],
        content_overrides: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, str]:
        skills_text = self._override_value(
            content_overrides,
            "skills",
            ", ".join([skill.skill_name for skill in candidate.skills]) or "N/A",
        )
        required = ", ".join(job.required_skills or []) if job else ""
        role = self._override_value(
            content_overrides,
            "role",
            job_title or (job.title if job else None) or candidate.main_role or "N/A",
        )
        name = self._override_value(content_overrides, "name", candidate.name or "N/A")
        summary = self._override_value(content_overrides, "summary", candidate.summary or "N/A")
        target_job = self._override_value(content_overrides, "target_job", job.title if job else role)
        phone = self._override_value(content_overrides, "phone", candidate.phone or "")
        email = self._override_value(content_overrides, "email", candidate.email or "")
        location = self._override_value(content_overrides, "location", candidate.location or "")
        work_lines = self._split_edit_lines(
            self._override_value(
                content_overrides,
                "work_experience",
                (
                    f"{role}\n"
                    f"{candidate.location or 'Remote'}\n"
                    "Relevant experience tailored for the target role.\n"
                    f"{skills_text}"
                ),
            )
        )
        education_lines = self._split_edit_lines(
            self._override_value(content_overrides, "education", "Available upon request")
        )
        project_lines = self._split_edit_lines(
            self._override_value(
                content_overrides,
                "projects",
                f"{target_job}\nPrepared for {target_job}. Required skills: {required or 'Not listed'}.",
            )
        )
        certification_lines = self._split_edit_lines(
            self._override_value(content_overrides, "certifications", "Available upon request")
        )
        language_lines = self._split_edit_lines(
            self._override_value(content_overrides, "languages", "English")
        )
        tool_lines = self._split_edit_lines(
            self._override_value(content_overrides, "tools", skills_text)
        )
        reference_lines = self._split_edit_lines(
            self._override_value(content_overrides, "references", "Available upon request.")
        )
        values = {
            "{{NAME}}": name,
            "{{ROLE}}": role,
            "{{SUMMARY}}": summary,
            "{{SKILLS}}": skills_text,
            "{{AVAILABILITY}}": candidate.availability or "Immediate",
            "{{EXPECTED_RATE}}": candidate.expected_rate or "TBD",
            "{{TARGET_JOB}}": target_job,
            "{{JOB_COMPANY}}": job.company if job else "",
            "{{JOB_PLATFORM}}": job.platform if job else "",
            "{{REQUIRED_SKILLS}}": required or "Not listed",
            "{{PHONE}}": phone,
            "{{EMAIL}}": email,
            "{{LOCATION}}": location,
            "[YOUR FULL NAME]": name,
            "[Job Title - e.g., Software Engineer / Accountant / Marketing Executive]": role,
            "[Job Title] - [Company Name]": self._line_or(work_lines, 0, role),
            "[Month Year - Month Year] | [City / Remote]": self._line_or(work_lines, 1, location or "Remote"),
            "[Achievement or responsibility with numbers/results if possible]": self._line_or(
                work_lines,
                2,
                "Relevant experience tailored for the target role.",
            ),
            "[Achievement or responsibility]": self._line_or(work_lines, 3, f"Target opportunity: {target_job}"),
            "[Tools/skills used]": self._line_or(work_lines, 4, skills_text),
            "[Degree / Diploma / Course Name]": self._line_or(education_lines, 0, "Available upon request"),
            "[Institution Name] | [Year]": self._line_or(education_lines, 1, ""),
            "[Relevant subjects, grade, or achievements]": self._line_or(education_lines, 2, ""),
            "[Project Name]": self._line_or(project_lines, 0, target_job),
            "[Short project description, your role, tools used, and result/link]": self._line_or(
                project_lines,
                1,
                f"Prepared for {target_job}. Required skills: {required or 'Not listed'}.",
            ),
            "[Certificate Name]": self._line_or(certification_lines, 0, "Available upon request"),
            "[Tool / Software 1]": self._line_or(tool_lines, 0, ""),
            "[Tool / Software 2]": self._line_or(tool_lines, 1, ""),
            "[Tool / Software 3]": self._line_or(tool_lines, 2, ""),
            "[Language 1 - Level]": self._line_or(language_lines, 0, "English"),
            "[Language 2 - Level]": self._line_or(language_lines, 1, ""),
            "Available upon request.": self._line_or(reference_lines, 0, "Available upon request."),
        }
        for index, skill in enumerate(self._split_edit_lines(skills_text)[:6], start=1):
            values[f"[Skill {index}]"] = skill
        values.update({key.lower(): value for key, value in values.items()})
        return values

    def _pdf_subtitle(self, candidate: Candidate, job: Optional[Job], template_file_name: str) -> str:
        role = job.title if job else (candidate.main_role or "Professional Profile")
        if template_file_name:
            return f"{role} - Format: {template_file_name}"
        return role
